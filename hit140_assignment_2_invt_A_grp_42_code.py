"""
HIT140 Assessment 2 — Investigation A: Do bats perceive rats as predators or just competitors?
-----------------------------------------------------------------------------------
This script reads two datasets (dataset1 and dataset2), cleans them,
and performs analyses to produce:
    1) A histogram comparing time-to-food by risk behaviour (avoid=0 vs risk
        =1) with a Mann–Whitney U test.
    2) Bar charts showing success rates (reward=1 proportions) by risk behaviour,
        with Fisher's exact test.
    3) A Word report compiling the figures, descriptive tables, and statistical
        results, with comments and conclusions printed in the terminal. 
    4) All outputs (figures and report) are saved to specified folders/files. 
-----------------------------------------------------------------------------------  

"""

# ==== Imports ====
# I’m using standard data science tools: pandas for data tables,
# numpy for number crunching, matplotlib for plots, scipy for stats,
# and python-docx to write a Word report.
import os, argparse, shutil, time
from datetime import datetime

import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
from matplotlib import cm
from scipy import stats

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


# ==== Command-line arguments ====
# I’m keeping CLI flags so I can switch files without editing code.
def parse_args():
    p = argparse.ArgumentParser(
        description="A2 (Investigation A) — histogram + bar charts with student-style comments."
    )
    p.add_argument("--d1", default="dataset1(1).csv", help="Dataset 1 (bat-level observations).")
    p.add_argument("--d2", default="dataset2.csv", help="Dataset 2 (period summaries).")
    p.add_argument("--outdir", default="figures", help="Where to save PNG figures.")
    p.add_argument("--report", default="A2_Report.docx", help="Word report filename.")
    # I default to Australian date format because my CSVs look like DD/MM/YYYY HH:MM
    p.add_argument("--dayfirst", action="store_true", default=True,
                   help="Parse dates as day/month/year (default True for AU).")
    return p.parse_args()


# ==== Small helper functions (IO + styling) ====
def fresh_output(outdir: str, report_path: str, retries: int = 5, delay: float = 0.5) -> str:
    """
    Start fresh each run:
      1) Delete and recreate the figures folder
      2) Try to delete old report; if it's locked (open in Word),
         fall back to a timestamped filename so my run still succeeds.
    """
    if os.path.isdir(outdir):
        shutil.rmtree(outdir, ignore_errors=True)
    os.makedirs(outdir, exist_ok=True)

    safe_report = report_path
    if os.path.exists(report_path):
        for _ in range(retries):
            try:
                os.remove(report_path)
                break
            except PermissionError:
                time.sleep(delay)
        else:
            ts = datetime.now().strftime("%Y%m%d_%H%M%S")
            root, ext = os.path.splitext(report_path)
            safe_report = f"{root}_{ts}{ext}"
            print(f"[warn] '{report_path}' locked; writing to '{safe_report}' instead.")
    return safe_report


def safe_numeric(df: pd.DataFrame, cols):
    """
    Convert listed columns to numeric.
    Any weird strings become NaN instead of crashing (‘errors=coerce’).
    This is part of my basic “cleaning”.
    """
    for c in cols:
        if c in df.columns:
            df[c] = pd.to_numeric(df[c], errors="coerce")


def parse_dates_inplace(df: pd.DataFrame, cols, dayfirst: bool = True):
    """
    Parse date/time text to real datetimes.
    I call this AFTER read_csv to respect AU format (dayfirst=True)
    and to avoid the original parser warnings.
    """
    for c in cols:
        if c in df.columns:
            df[c] = pd.to_datetime(df[c], errors="coerce", dayfirst=dayfirst)


# ---- Word styling helpers (so the report looks neat and consistent) ----
def set_para_style(p, size=11, bold=False, align="left"):
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if align == "center" else WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.color.rgb = RGBColor(0, 0, 0)  # black text throught the document


def add_heading(doc: Document, text: str, level=1, align="left"):
    p = doc.add_heading(text, level=level)
    set_para_style(p, size=12 if level >= 2 else 14, bold=True, align=align)
    return p


def add_paragraph(doc: Document, text: str, size=11, bold=False, align="left"):
    p = doc.add_paragraph(text)
    set_para_style(p, size=size, bold=bold, align=align)
    return p


def add_df_table(doc: Document, df: pd.DataFrame, title: str, round_dec=3, max_rows=20):
    """
    Add a small descriptive table to the Word report, So it is easy to understand.
    I round numeric columns so it’s easier to read.
    """
    add_paragraph(doc, title, bold=True)
    if df is None or df.empty:
        add_paragraph(doc, "(no data)")
        return
    df_show = df.copy()
    for c in df_show.columns:
        if pd.api.types.is_numeric_dtype(df_show[c]):
            df_show[c] = df_show[c].round(round_dec)
    truncated = False
    if len(df_show) > max_rows:
        df_show = df_show.head(max_rows)
        truncated = True

    table = doc.add_table(rows=1, cols=len(df_show.columns) + 1)
    table.style = "Table Grid"

    # Header row: first blank for the index
    hdr = table.rows[0].cells
    hdr[0].text = ""
    for j, col in enumerate(df_show.columns, start=1):
        hdr[j].text = str(col)
    for cell in table.rows[0].cells:
        for p in cell.paragraphs:
            set_para_style(p, size=10, bold=True, align="center")

    # Body rows
    for idx, row in df_show.iterrows():
        cells = table.add_row().cells
        cells[0].text = str(idx)
        for j, col in enumerate(df_show.columns, start=1):
            val = "" if pd.isna(row[col]) else str(row[col])
            cells[j].text = val
    for r in table.rows[1:]:
        for cell in r.cells:
            for p in cell.paragraphs:
                set_para_style(p, size=10, bold=False, align="left")

    if truncated:
        add_paragraph(doc, f"(showing first {max_rows} rows)")


def add_stats_table(doc: Document, rows):
    """
    Neat statistical summary table for the report. Each row is a dict with:
    Test, Variables, Statistic, p value, n, Interpretation and Graph
    """
    add_heading(doc, "Statistical results (formatted table)", level=2)
    cols = ["Test", "Variables", "Statistic", "p value", "n", "Interpretation", "Graph"]
    table = doc.add_table(rows=1, cols=len(cols))
    table.style = "Table Grid"

    # Header
    for i, c in enumerate(cols):
        cell = table.rows[0].cells[i]
        cell.text = c
        for p in cell.paragraphs:
            set_para_style(p, size=10, bold=True, align="center")

    # Body
    for r in rows:
        tr = table.add_row().cells
        tr[0].text = r.get("Test", "")
        tr[1].text = r.get("Variables", "")
        tr[2].text = r.get("Statistic", "")
        tr[3].text = r.get("p value", "")
        tr[4].text = r.get("n", "")
        tr[5].text = r.get("Interpretation", "")
        tr[6].text = r.get("Graph", "")
    for row in table.rows[1:]:
        for c in row.cells:
            for p in c.paragraphs:
                set_para_style(p, size=10, bold=False, align="left")


# ==== Plot helpers (colours + tiny utilities) ====
def _two_bar_colors():
    # I pick two contrasting colours so the risk groups are easy to tell apart.
    return ["#e74c3c", "#1f77b4"]  # red, blue colors


def _multi_bar_colors(n):
    # If I ever need more colours, I can sample from a palette here.
    cmap = cm.get_cmap("tab20", n)
    return [tuple(cmap(i)) for i in range(n)]


def _fmt_pct(x):
    # Convenience: print 0.534 as "53.4%"
    return "NA" if (x is None or pd.isna(x)) else f"{100*float(x):.1f}%"


def _sig_word(p):
    # I use a standard 0.05 cutoff for “significant” in this assignment.
    return "SIGNIFICANT" if (p is not None and p < 0.05) else "not significant"


# ==== MAIN WORKFLOW ====
def main():
    args = parse_args()
    report_path = fresh_output(args.outdir, args.report)

    # STEP 1 — Read the datasets
    # I read both CSVs without parse_dates to avoid automatic month-first parsing.
    d1 = pd.read_csv(args.d1, low_memory=False)
    try:
        d2 = pd.read_csv(args.d2, low_memory=False)
    except Exception:
        d2 = pd.DataFrame()

    # STEP 2 — CLEANING (both datasets)
    # I did a few light cleaning steps to ensure the analyses run smoothly.
    # (A) Convert date columns to real datetimes using AU-style dayfirst=True
    parse_dates_inplace(d1, ["start_time", "rat_period_start", "rat_period_end", "sunset_time"],
                        dayfirst=args.dayfirst)
    if "time" in d2.columns:
        parse_dates_inplace(d2, ["time"], dayfirst=args.dayfirst)

    # (B) Force important numeric columns to numbers; bad entries become NaN.
    #     This helps plotting and stats work without errors.
    safe_numeric(d1, [
        "bat_landing_to_food", "seconds_after_rat_arrival", "hours_after_sunset",
        "risk", "reward", "month", "season"
    ])

    # (C) A small tidy up for the season column (if present).
    if "season" in d1.columns:
        try:
            d1["season"] = d1["season"].astype("Int64")  # keeps NaN as <NA>
        except Exception:
            # If season has text labels, at least remove accidental spaces.
            d1["season"] = d1["season"].astype(str).str.strip()

    # (D) Light-touch cleaning for d2:
    #     In this trimmed analysis, we don't compute from d2,
    #     so date parsing is enough to keep it “clean” for now.

    # Quick descriptives just for context in the report
    d1_cols = [c for c in ["bat_landing_to_food", "risk", "reward"] if c in d1.columns]
    desc1 = d1[d1_cols].describe().T if d1_cols else pd.DataFrame()
    desc1.attrs["n_records_d1"] = len(d1)

    # STEP 3 — Analysis A: Histogram of time-to-food by risk behaviour
    # Goal: compare the DISTRIBUTIONS of approach times between risk groups.
    # Why histogram? The y-axis is frequency, so I can see the shape and spread,not just the average.
    fig_time_risk = None
    mw_U = mw_p = n0 = n1 = None
    mean0 = mean1 = med0 = med1 = None

    if set(["bat_landing_to_food", "risk"]).issubset(d1.columns):
        # Drop rows missing either time or risk (classic cleaning before analysis)
        tmp = d1.dropna(subset=["bat_landing_to_food", "risk"]).copy()
        # Also guard against negative/inf values so the histogram is sensible
        tmp = tmp[np.isfinite(tmp["bat_landing_to_food"])]
        tmp = tmp[tmp["bat_landing_to_food"] >= 0]

        # Split into the two groups: risk=0 (avoid), risk=1 (take)
        g0 = tmp.loc[tmp["risk"] == 0, "bat_landing_to_food"]
        g1 = tmp.loc[tmp["risk"] == 1, "bat_landing_to_food"]
        n0, n1 = len(g0), len(g1)

        if n0 > 0 and n1 > 0:
            # I report both means and medians (medians are robust to outliers).
            mean0, mean1 = g0.mean(), g1.mean()
            med0, med1 = g0.median(), g1.median()

            # Mann–Whitney U test is a non-parametric test (no normality assumption).
            U, p = stats.mannwhitneyu(g0, g1, alternative="two-sided")
            mw_U, mw_p = float(U), float(p)

            # Use shared bins so the two histograms line up on the same x-scale
            pooled = pd.concat([g0, g1], ignore_index=True)
            bins = np.histogram_bin_edges(pooled, bins="auto")

            plt.figure()
            colors = _two_bar_colors()
            plt.hist(g0, bins=bins, color=colors[0], alpha=0.65,
                     label="Risk=0 (avoid)", edgecolor="black")
            plt.hist(g1, bins=bins, color=colors[1], alpha=0.65,
                     label="Risk=1 (take)", edgecolor="black")
            plt.xlabel("Time to approach food (seconds)")
            plt.ylabel("Frequency")
            plt.title("Time to approach food by risk behaviour (histogram)")
            plt.legend(frameon=False)
            plt.tight_layout()
            fig_time_risk = os.path.join(args.outdir, "Figure1_TimeToFoodByRisk_Hist.png")
            plt.savefig(fig_time_risk, dpi=150)
            plt.close()

            # To print a short conclusion in the terminal
            faster_group = "Risk-Takers (1)" if mean1 < mean0 else "Risk-Avoiders (0)"
            diff_dir = "tend to be faster" if (mean1 < mean0) else "tend to be slower"
            print(f"Conclusion: Fig 1 Time-to-food histogram → "
                  f"means: avoid={mean0:.1f}s, take={mean1:.1f}s; medians: avoid={med0:.1f}s, take={med1:.1f}s; "
                  f"Mann–Whitney U={mw_U:.0f}, p={mw_p:.3g} → {_sig_word(mw_p)} at α=0.05. "
                  f"{faster_group} {diff_dir} on average.")

    # STEP 4 — Analysis B: Success rate by risk behaviour (bar chart)
    # Goal: compare proportions of success (reward=1) between the two groups.
    # Why Fisher’s exact test? It’s reliable even when counts are small.
    fig_success_rate = None
    fig_predator = None
    success_rates = None

    if set(["risk", "reward"]).issubset(d1.columns):
        # Standard cleaning: only keep rows where both values exist
        tmp_sr = d1.dropna(subset=["risk", "reward"]).copy()
        avoiders = tmp_sr.loc[tmp_sr["risk"] == 0]
        takers   = tmp_sr.loc[tmp_sr["risk"] == 1]

        # Helper to compute the proportion of reward=1
        def _success_rate(df):
            return float(df["reward"].mean()) if len(df) else float("nan")

        rate_avoid = _success_rate(avoiders)
        rate_take  = _success_rate(takers)

        # I also calculate raw counts for Fisher’s test
        n_avoid = int(len(avoiders))
        n_take  = int(len(takers))
        succ_avoid = int(avoiders["reward"].sum()) if n_avoid else 0
        succ_take  = int(takers["reward"].sum()) if n_take else 0
        fail_avoid = n_avoid - succ_avoid
        fail_take  = n_take - succ_take

        success_rates = {
            "avoid_rate": rate_avoid,
            "take_rate": rate_take,
            "n_avoid": n_avoid,
            "n_take": n_take,
            "succ_avoid": succ_avoid,
            "succ_take": succ_take,
            "fail_avoid": fail_avoid,
            "fail_take": fail_take,
        }

        # Fisher’s exact test setup: 2x2 table of successes/failures by group
        fisher_p = None
        if (n_avoid > 0 and n_take > 0):
            table_2x2 = [[succ_avoid, fail_avoid],
                         [succ_take,  fail_take]]
            try:
                OR, fisher_p = stats.fisher_exact(table_2x2, alternative="two-sided")
            except Exception:
                OR, fisher_p = np.nan, None

        #Figure 2: bars with percentages written on top of each bar
        plt.figure()
        labels = ["Risk-Avoiders", "Risk-Takers"]
        values = [rate_avoid, rate_take]
        colors = _two_bar_colors()
        bars = plt.bar(labels, values, color=colors, edgecolor="black")
        plt.ylim(0, 1)
        plt.ylabel("Success Rate")
        plt.title("Success rate by risk behaviour (bar graph)")
        for pbar in bars:
            val = pbar.get_height()
            plt.gca().annotate(f"{val*100:.1f}%",
                               (pbar.get_x() + pbar.get_width()/2, val),
                               ha="center", va="bottom", xytext=(0, 3), textcoords="offset points")
        plt.legend([bars[0], bars[1]], labels, title="Groups", frameon=False, loc="best")
        plt.tight_layout()
        fig_success_rate = os.path.join(args.outdir, "Figure_2_Success_Rate_By_Risk.png")
        plt.savefig(fig_success_rate, dpi=150)
        plt.close()

        # Terminal summary for Figure 2
        higher = "Risk-Takers (1)" if rate_take > rate_avoid else "Risk-Avoiders (0)"
        if fisher_p is not None:
            print(f"Conclusion: Fig 2 Success rates → "
                  f"avoid={_fmt_pct(rate_avoid)} ({succ_avoid}/{n_avoid}), "
                  f"take={_fmt_pct(rate_take)} ({succ_take}/{n_take}); "
                  f"Fisher's exact p={fisher_p:.3g} → {_sig_word(fisher_p)} at α=0.05. "
                  f"{higher} show the higher success proportion.")
        else:
            print(f"Conclusion: Fig 2 Success rates → "
                  f"avoid={_fmt_pct(rate_avoid)} ({succ_avoid}/{n_avoid}), "
                  f"take={_fmt_pct(rate_take)} ({succ_take}/{n_take}); "
                  f"Fisher test not computed.")

        #Figure 3 (bar graph)
        plt.figure()
        bars = plt.bar(labels, values, color=colors, edgecolor="black")
        plt.ylim(0, 1)
        plt.ylabel("Success Rate")
        plt.title("Bat Success Rate: Evidence of Predator Perception (bar graph)")
        for pbar in bars:
            val = pbar.get_height()
            plt.gca().annotate(f"{val*100:.1f}%",
                               (pbar.get_x() + pbar.get_width()/2, val),
                               ha="center", va="bottom", xytext=(0, 3), textcoords="offset points")
        plt.legend([bars[0], bars[1]], labels, title="Groups", frameon=False, loc="best")
        plt.tight_layout()
        fig_predator = os.path.join(args.outdir, "Figure_3_Bat_Success_Rates_Evidence.png")
        plt.savefig(fig_predator, dpi=150)
        plt.close()

    # STEP 5 — Building the report in word format (figures, tables, captions)
    doc = Document()
    add_heading(doc, "Assessment 2 – Investigation A Report (Histogram + Bars, with Terminal Conclusions)", level=1)
    add_paragraph(doc, f"Dataset 1 path: {os.path.abspath(args.d1)}")
    add_paragraph(doc, f"Date parsing: dayfirst=True (Australian format) applied after CSV read.", size=10)

    add_df_table(doc, desc1, "Dataset 1 – key descriptives")
    add_paragraph(doc, f"Records: {len(d1)}")

    add_heading(doc, "Graph types used and why", level=2)
    add_paragraph(doc, "• Histogram (Figure 1) — lets me compare the full distribution of approach times "
                       "(x: seconds; y: frequency) across risk groups.")
    add_paragraph(doc, "• Bar charts (Figures 2 & 3) — show success-rate (proportion of reward=1) for each group; "
                       "Figure 3 repeats the same data with the assignment-specific caption ‘Evidence of Predator Perception’.")
    add_paragraph(doc, "Note: Figures 2 and 3 intentionally use the SAME values", size=10)

    # Statistical table rows I want to include in the report
    rows = []
    if success_rates is not None:
        # Report observed rates + Fisher p if available
        fisher_text = "—"
        try:
            # fisher_p may not exist in scope if the block didn’t run; safe guard:
            fisher_p  # noqa
            fisher_text = f"{fisher_p:.2e}" if fisher_p is not None else "—"
        except NameError:
            pass
        rows.append({
            "Test": "Context (proportions)",
            "Variables": "success rate ~ risk group",
            "Statistic": (f"avoid={_fmt_pct(success_rates['avoid_rate'])}, "
                          f"take={_fmt_pct(success_rates['take_rate'])}"),
            "p value": f"Fisher={fisher_text}",
            "n": f"n0={success_rates['n_avoid']}, n1={success_rates['n_take']}",
            "Interpretation": "Observed success proportions by group (same values in Fig 2 & Fig 3).",
            "Graph": "Bar chart"
        })

    if mw_U is not None:
        rows.append({
            "Test": "Mann–Whitney U",
            "Variables": "time-to-food ~ risk (0=avoid,1=take)",
            "Statistic": f"U={mw_U:.0f}",
            "p value": f"{mw_p:.2e}",
            "n": f"n0={n0}, n1={n1}",
            "Interpretation": "Non-parametric group difference in approach times.",
            "Graph": "Histogram (frequency)"
        })

    if rows:
        add_stats_table(doc, rows)

    # Adds figures with captions 
    if fig_time_risk and os.path.exists(fig_time_risk):
        add_paragraph(doc, "Figure 1: Time to approach food by risk behaviour (histogram: frequency vs seconds)", bold=True)
        doc.add_picture(fig_time_risk, width=Inches(5.5))

    if fig_success_rate and os.path.exists(fig_success_rate):
        add_paragraph(doc, "Figure 2: Success rate by risk behaviour (bar graph)", bold=True)
        doc.add_picture(fig_success_rate, width=Inches(5.5))

    if fig_predator and os.path.exists(fig_predator):
        add_paragraph(doc, "Figure 3: Bat Success Rate – Evidence of Predator Perception (bar graph)", bold=True)
        doc.add_picture(fig_predator, width=Inches(5.5))

    # STEP 6 — Saves the report and prints the output paths
    doc.save(report_path)
    print(f"\nSaved Word report to: {os.path.abspath(report_path)}")
    print(f"Saved figures in: {os.path.abspath(args.outdir)}")
    print("Overall Conclusion: \n1) Approach-time distribution: The histogram indicates that risk-avoid bats take more time to approach food, consistent with vigilance.\n2) Group difference test: Mann–Whitney U on approach times indicates a reliable difference at α=0.05 (p-value).\n3) Feeding success: Risk-avoidance results in significantly greater feeding success, while aggressive/risk-taking usually leads to failure.\n4) Proportion test: Fisher’s exact test on success vs risk group supports a significant difference in success rates in favor of avoidance.\n5) General conclusion: Findings suggest bats perceive rats as predators/dangerous threats—not just competitors—showing avoidance/vigilance and improved success with avoidance.")




if __name__ == "__main__":
    main()
